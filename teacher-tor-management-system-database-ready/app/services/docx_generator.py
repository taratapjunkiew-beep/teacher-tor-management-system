import json
import re
import uuid
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_ORIENTATION, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from ..config import GENERATED_DIR
from ..models import Course

FONT_NAME = "TH SarabunPSK"
BASE_SIZE = 16
TEMPLATE_DIR = Path(__file__).resolve().parents[1] / "assets"
EMBLEM_PATH = TEMPLATE_DIR / "ovec_emblem.jpeg"

# ---------- low-level formatting ----------
def _font(run, size: float = BASE_SIZE, bold: bool = False):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold


def _set_doc_defaults(doc: Document):
    for style_name in ["Normal", "Header", "Footer"]:
        try:
            style = doc.styles[style_name]
            style.font.name = FONT_NAME
            style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
            style._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
            style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
            style.font.size = Pt(BASE_SIZE)
        except Exception:
            pass


def set_portrait(section):
    section.orientation = WD_ORIENTATION.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(1.8)


def set_landscape(section):
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)


def add_p(doc: Document, text: str = "", size: float = BASE_SIZE, bold: bool = False,
          align=None, first_line: float | None = None, before: float = 0, after: float = 0,
          line_spacing: float = 1.0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line_spacing
    if first_line is not None:
        p.paragraph_format.first_line_indent = Cm(first_line)
    if align is not None:
        p.alignment = align
    r = p.add_run(str(text or ""))
    _font(r, size, bold)
    return p


def add_multiline(doc: Document, text: str, *, size=BASE_SIZE, bold=False, align=None, first_line=None):
    for para in split_paragraphs(text):
        add_p(doc, para, size=size, bold=bold, align=align, first_line=first_line)


def split_paragraphs(text: str) -> list[str]:
    text = str(text or "").replace("\r", "").strip()
    if not text:
        return [""]
    # keep explicit paragraph breaks but collapse accidental whitespace
    parts = [re.sub(r"[ \t]+", " ", p).strip() for p in re.split(r"\n\s*\n|\n", text) if p.strip()]
    return parts or [""]


def set_cell_text(cell, text: str = "", size: float = BASE_SIZE, bold: bool = False,
                  align=WD_ALIGN_PARAGRAPH.CENTER, fill: str | None = None, first_line: float | None = None):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    if first_line is not None:
        p.paragraph_format.first_line_indent = Cm(first_line)
    lines = str(text if text is not None else "").split("\n")
    if not lines:
        lines = [""]
    for idx, line in enumerate(lines):
        if idx:
            p.add_run().add_break()
        r = p.add_run(str(line))
        _font(r, size, bold)
    if fill:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)


def shade_cell(cell, fill="D9EAF7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_cm: float):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_grid(table, widths: list[float]):
    tbl = table._tbl
    # remove existing grid and write explicit column grid
    old = tbl.tblGrid
    if old is not None:
        tbl.remove(old)
    grid = OxmlElement("w:tblGrid")
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(w * 567)))
        grid.append(col)
    tbl.insert(0, grid)


def style_table(table, size: float = BASE_SIZE, widths: list[float] | None = None):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    # Force Word/LibreOffice to respect column widths instead of autofitting tiny columns.
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    if widths:
        set_table_grid(table, widths)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if widths and idx < len(widths):
                set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    _font(run, size, run.bold)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


# ---------- data helpers ----------
def jlist(raw) -> list:
    try:
        data = json.loads(raw or "[]")
        return data if isinstance(data, list) else []
    except Exception:
        return []


def clean_text(s: str) -> str:
    s = str(s or "").replace("\r", "")
    s = re.sub(r"[ \t]+", " ", s)
    s = s.replace(" สาขาวิชา สาขาวิชา ", " สาขาวิชา ")
    return s.strip()


def clean_course_description(text: str) -> str:
    """Remove footer/header fragments accidentally captured from PDFs."""
    text = clean_text(text)
    # Some extracted PDFs append curriculum/footer text after the real course description.
    text = re.split(r"\n?หลักสูตรประกาศนียบัตร|\n?หลักสูตรรายวิชา|\n?สาขาวิชา(?=เทคโนโลยี|อิเล็กทรอนิกส์|ช่าง|เมคคาทรอนิกส์)", text, maxsplit=1)[0].strip()
    text = re.sub(r"ปประเภทวิชา", "ประเภทวิชา", text)
    return text


def short_curriculum(course: Course) -> str:
    cur = clean_text(course.curriculum)
    if "ชั้นสูง" in cur or course.level == "ปวส.":
        return "ประกาศนียบัตรวิชาชีพชั้นสูง (ปวส.)"
    return "ประกาศนียบัตรวิชาชีพ (ปวช.)"


def is_generic_unit(title: str, course_name: str) -> bool:
    t = str(title or "")
    generic_terms = [
        "หลักการพื้นฐานของ", "เครื่องมือและอุปกรณ์ที่ใช้ใน", "การปฏิบัติงานพื้นฐานใน",
        "การวิเคราะห์และแก้ไขปัญหาใน", "การทดสอบและประเมินผลงาน", "โครงงานหรือภาระงานใน",
        "การประยุกต์ใช้"
    ]
    return any(x in t for x in generic_terms) and course_name in t


def build_units_from_description(course: Course) -> list[dict[str, Any]]:
    existing = jlist(course.units)
    if existing and len(existing) >= 8 and not all(is_generic_unit(u.get("title", ""), course.name) for u in existing[:6]):
        return existing[:10]

    desc = clean_course_description(course.description)
    desc = re.sub(r"^ศึกษาและปฏิบัติ(งาน)?เกี่ยวกับ", "", desc)
    desc = re.sub(r"^ศึกษาเกี่ยวกับ", "", desc)
    # split descriptive phrase into usable topic chunks
    raw = re.split(r",| และ | การ(?=[ก-ฮ])", desc)
    topics = []
    for item in raw:
        t = clean_text(item.strip(" .;:、"))
        if len(t) < 8:
            continue
        # strip very long trailing text and normalize common starts
        t = re.sub(r"^(เกี่ยวกับ|เรื่อง)", "", t).strip()
        if t and t not in topics:
            topics.append(t)

    # build better titles from actual description topics
    titles = []
    for t in topics:
        # keep unit titles concise so tables do not become cramped
        t = re.sub(r"^(วัดทดสอบวงจร|วัดทดสอบ)", "การวัดและทดสอบ", t)
        t = re.sub(r"^(ประกอบและทดสอบ)", "การประกอบและทดสอบ", t)
        if len(t) > 54:
            t = t[:54].rsplit(" ", 1)[0]
        titles.append(t)
    if len(titles) < 10:
        fallback = [
            f"หลักการและภาพรวมของ{course.name}",
            f"เครื่องมือ วัสดุ และอุปกรณ์ใน{course.name}",
            f"การอ่านแบบและวิเคราะห์ข้อกำหนดของ{course.name}",
            f"การเตรียมและวางแผนการปฏิบัติงาน{course.name}",
            f"การปฏิบัติงานตามขั้นตอนของ{course.name}",
            f"การทดสอบ ตรวจสอบ และบันทึกผล{course.name}",
            f"การวิเคราะห์ปัญหาและแก้ไขข้อบกพร่องใน{course.name}",
            f"การประยุกต์ใช้{course.name}ในงานอาชีพ",
            f"การประเมินผลงานและนำเสนอ{course.name}",
            f"โครงการหรือภาระงานบูรณาการ{course.name}",
        ]
        for f in fallback:
            if f not in titles:
                titles.append(f)
            if len(titles) >= 10:
                break
    titles = titles[:10]

    total_t, total_p = target_hours(course)
    theory_parts = distribute(total_t, 10)
    practice_parts = distribute(total_p, 10)
    units = []
    for i, title in enumerate(titles, 1):
        bloom = bloom_for_unit(i)
        units.append({
            "no": i,
            "title": title,
            "theory": theory_parts[i-1],
            "practice": practice_parts[i-1],
            "k": k_for_unit(i),
            "s": "S5" if i == 10 else "S4",
            "a": "A5",
            "ap": "AP5" if i == 10 else "AP4" if course.level == "ปวส." else "AP3",
            "bloom": bloom,
        })
    return units


def target_hours(course: Course) -> tuple[int, int]:
    # Prefer existing unit sums imported from database when available.
    existing = jlist(course.units)
    if existing:
        total_t = sum(int(x.get("theory", 0) or 0) for x in existing)
        total_p = sum(int(x.get("practice", 0) or 0) for x in existing)
        if total_t > 0 or total_p > 0:
            return total_t, total_p
    # fallback: ปวช. often uses 18/54 style; ปวส. templates may vary, but keep safe standard.
    if course.level == "ปวช.":
        return int(course.theory_hours) * 18, int(course.practice_hours) * 18
    return int(course.theory_hours) * 15, int(course.practice_hours) * 15


def distribute(total: int, n: int) -> list[int]:
    base = total // n
    rem = total - base * n
    return [base + (1 if i < rem else 0) for i in range(n)]


def bloom_for_unit(i: int) -> list[str]:
    # six cognitive dimensions: knowledge, understand, apply, analyze, evaluate, create
    patterns = {
        1: ["-", "20", "-", "20", "-", "-"],
        2: ["-", "20", "-", "20", "-", "-"],
        3: ["-", "20", "20", "-", "-", "-"],
        4: ["-", "20", "20", "-", "-", "-"],
        5: ["-", "20", "20", "-", "-", "-"],
        6: ["-", "20", "-", "20", "-", "-"],
        7: ["-", "20", "-", "-", "-", "20"],
        8: ["-", "20", "-", "20", "-", "-"],
        9: ["-", "20", "-", "-", "20", "-"],
        10: ["-", "-", "20", "-", "-", "20"],
    }
    return patterns.get(i, ["-", "20", "20", "-", "-", "-"])


def k_for_unit(i: int) -> str:
    mapping = {
        1: "K2, K4", 2: "K2, K4", 3: "K2, K3", 4: "K2, K3", 5: "K2, K3",
        6: "K2, K4", 7: "K2, K6", 8: "K2, K4", 9: "K2, K5", 10: "K3, K6",
    }
    return mapping.get(i, "K2, K3")


def normalize_ap(ap: str) -> str:
    return str(ap or "AP3").replace("Ap", "AP")


def competencies(course: Course) -> list[str]:
    data = [clean_text(x) for x in jlist(course.competencies) if clean_text(x)]
    return data or ["ปฏิบัติงานตามสมรรถนะรายวิชา"]


def pick_comp(course: Course, idx: int) -> str:
    comps = competencies(course)
    return comps[min(idx, len(comps)-1)]


# ---------- content builders ----------
def add_cover(doc: Document, course: Course, teacher: dict[str, Any]):
    if EMBLEM_PATH.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run()
        r.add_picture(str(EMBLEM_PATH), width=Cm(3.4))
    else:
        add_p(doc, "", 16, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_p(doc, "โครงการสอน", 24, True, WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_p(doc, f"หลักสูตร{short_curriculum(course)} พุทธศักราช 2567", 20, True, WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, f"ประเภทวิชา {course.program_type}", 20, True, WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, f"กลุ่มอาชีพ {course.occupational_group}", 20, True, WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, f"สาขาวิชา {course.major}", 20, True, WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "", 16)
    add_p(doc, f"รหัสวิชา {course.code}", 22, True, WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, f"วิชา {course.name}", 22, True, WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(6):
        add_p(doc, "", 16)
    add_p(doc, "จัดทำโดย", 18, True, WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, teacher.get("teacher_name", ""), 18, True, WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, f"ตำแหน่ง {teacher.get('position', '')}", 18, False, WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(3):
        add_p(doc, "", 16)
    add_p(doc, teacher.get("college", "วิทยาลัยการอาชีพบ้านผือ"), 18, True, WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "สำนักงานคณะกรรมการการอาชีวศึกษา", 18, False, WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "กระทรวงศึกษาธิการ", 18, False, WD_ALIGN_PARAGRAPH.CENTER)


def add_course_syllabus(doc: Document, course: Course):
    add_p(doc, "หลักสูตรรายวิชา", 20, True, WD_ALIGN_PARAGRAPH.CENTER, after=6)
    add_p(doc, f"หลักสูตร {short_curriculum(course)}", 16, False, WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, f"ประเภทวิชา {course.program_type}", 16, False, WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, f"กลุ่มอาชีพ {course.occupational_group}", 16, False, WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, f"สาขาวิชา {course.major}", 16, False, WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, f"รหัสวิชา {course.code} ชื่อวิชา {course.name}", 16, False, WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, f"ทฤษฎี {course.theory_hours} ชั่วโมง/สัปดาห์  ปฏิบัติ {course.practice_hours} ชั่วโมง/สัปดาห์   จำนวน {course.credits} หน่วยกิต", 16, False, WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "", 16)

    sections = [
        ("อ้างอิงมาตรฐาน", clean_text(course.standard_ref) or "-"),
        ("ผลลัพธ์การเรียนรู้ระดับรายวิชา", clean_text(course.learning_outcome)),
    ]
    for heading, body in sections:
        add_p(doc, heading, 16, True, WD_ALIGN_PARAGRAPH.LEFT, after=0)
        add_multiline(doc, body, first_line=1.0)
        add_p(doc, "", 8)

    add_p(doc, "จุดประสงค์รายวิชา เพื่อให้", 16, True)
    for idx, obj in enumerate(jlist(course.objectives), 1):
        add_p(doc, f"{idx}. {clean_text(obj)}", 16, False, first_line=0)
    add_p(doc, "", 8)

    add_p(doc, "สมรรถนะรายวิชา", 16, True)
    for idx, comp in enumerate(competencies(course), 1):
        add_p(doc, f"{idx}. {comp}", 16, False, first_line=0)
    add_p(doc, "", 8)

    add_p(doc, "คำอธิบายรายวิชา", 16, True)
    add_multiline(doc, clean_course_description(course.description), first_line=1.0)


def add_standard_table(doc: Document, course: Course):
    add_p(doc, "มาตรฐานอาชีพ", 20, True, WD_ALIGN_PARAGRAPH.CENTER, after=6)
    std_text = clean_text(course.occupational_standard or course.standard_ref or "-")
    if std_text and std_text != "-":
        add_multiline(doc, std_text, first_line=1.0)
    else:
        add_p(doc, "-", 16, False, WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "", 8)

    widths = [2.1, 4.5, 2.3, 4.8, 5.3, 5.0, 2.6, 2.3]
    t = doc.add_table(rows=2, cols=8)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Main header with merged cells like the model file.
    row0 = t.rows[0].cells
    row0[0].merge(row0[1]); set_cell_text(row0[0], "หน่วยสมรรถนะ", 14, True, fill="D9EAF7")
    row0[2].merge(row0[3]); set_cell_text(row0[2], "สมรรถนะย่อย", 14, True, fill="D9EAF7")
    set_cell_text(row0[4], "เกณฑ์การปฏิบัติงาน", 14, True, fill="D9EAF7")
    set_cell_text(row0[5], "วิธีประเมิน", 14, True, fill="D9EAF7")
    set_cell_text(row0[6], "รหัส PC\n(ตามเล่มมาตรฐาน)", 14, True, fill="D9EAF7")
    set_cell_text(row0[7], "รหัส PC\n(จากระบบ)", 14, True, fill="D9EAF7")
    row1 = t.rows[1].cells
    for i, h in enumerate(["รหัส (1.1)", "คำอธิบาย (1.2)", "รหัส (2.1)", "คำอธิบาย (2.2)", "", "", "", ""]):
        set_cell_text(row1[i], h, 14, True, fill="D9EAF7")
    for i in range(4, 8):
        row0[i].merge(row1[i])
    set_repeat_table_header(t.rows[0])
    set_repeat_table_header(t.rows[1])

    comps = competencies(course)
    # Do not invent PC codes. Only use simple structure and mark unavailable PC with '-'.
    codes = re.findall(r"\b\d{5}\.\d{2}\b|\b\d{5}\.\d{2}\.\d{2}\b", std_text)
    if not codes:
        codes = ["-"] * min(3, max(1, len(comps)))
    rows = []
    for i, comp in enumerate(comps[:3]):
        unit_code = codes[i] if i < len(codes) else "-"
        rows.append([
            unit_code,
            "อาชีพ/หน่วยสมรรถนะอ้างอิงตามหลักสูตร" if unit_code == "-" else "หน่วยสมรรถนะตามอ้างอิงมาตรฐาน",
            "-",
            comp,
            "ปฏิบัติงานตามแบบ/เงื่อนไขที่กำหนดและสอดคล้องกับสมรรถนะรายวิชา",
            "ประเมินจากชิ้นงาน การสาธิต แบบทดสอบ รายงาน และการนำเสนอ",
            "-",
            "-",
        ])
    for data in rows:
        r = t.add_row().cells
        for i, value in enumerate(data):
            set_cell_text(r[i], value, 14, False, WD_ALIGN_PARAGRAPH.LEFT if i in [1,3,4,5] else WD_ALIGN_PARAGRAPH.CENTER)
    style_table(t, 14, widths)
    add_p(doc, "หมายเหตุ: ไม่สร้างรหัสหน่วยสมรรถนะ/รหัส PC เพิ่มเติม หากไม่มีในเอกสารต้นฉบับ", 16, False)


def add_learning_unit_analysis(doc: Document, course: Course, units: list[dict[str, Any]]):
    add_p(doc, "ตารางวิเคราะห์หน่วยการเรียนรู้", 20, True, WD_ALIGN_PARAGRAPH.CENTER, after=6)
    widths = [4.7, 4.7, 5.2, 5.0, 5.0]
    t = doc.add_table(rows=2, cols=5)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    merged = t.rows[0].cells[0].merge(t.rows[0].cells[4])
    set_cell_text(merged, "ผลลัพธ์การเรียนรู้ระดับรายวิชา:\n" + clean_text(course.learning_outcome), 16, True, WD_ALIGN_PARAGRAPH.LEFT, fill="F3F7FB")
    headers = ["งานหลัก (Duty)", "งานย่อย (Task)", "สมรรถนะย่อย\n(มาตรฐานอาชีพ)", "ความรู้ในการปฏิบัติงาน", "ทักษะในการปฏิบัติงาน"]
    for i, h in enumerate(headers):
        set_cell_text(t.rows[1].cells[i], h, 16, True, fill="D9EAF7")
    set_repeat_table_header(t.rows[0])
    set_repeat_table_header(t.rows[1])

    for i, item in enumerate(units):
        r = t.add_row().cells
        duty = item["title"] if i == 0 or i % 2 == 0 else ""
        set_cell_text(r[0], duty, 16, False, WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(r[1], f"ศึกษา/ฝึกปฏิบัติ {item['title']}", 16, False, WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(r[2], pick_comp(course, i), 16, False, WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(r[3], f"ความรู้เกี่ยวกับ{item['title']}", 16, False, WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(r[4], f"ทักษะการปฏิบัติงานเกี่ยวกับ{item['title']}", 16, False, WD_ALIGN_PARAGRAPH.LEFT)
    style_table(t, 16, widths)


def add_behavior_table1(doc: Document, course: Course, units: list[dict[str, Any]]):
    add_p(doc, "ตารางวิเคราะห์พฤติกรรมการเรียนรู้", 20, True, WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_course_line(doc, course)
    widths = [7.8, 1.2, 1.25, 1.2, 1.25, 1.25, 1.25, 1.3, 1.3, 1.3, 1.6, 1.55, 1.7]
    t = doc.add_table(rows=3, cols=13)
    # Top header
    r0 = t.rows[0].cells
    r0[0].merge(t.rows[2].cells[0]); set_cell_text(r0[0], "หน่วยการเรียนรู้", 14, True, fill="D9EAF7")
    r0[1].merge(r0[9]); set_cell_text(r0[1], "ระดับความสามารถที่คาดหวัง", 14, True, fill="D9EAF7")
    r0[10].merge(t.rows[2].cells[10]); set_cell_text(r0[10], "รวม", 14, True, fill="D9EAF7")
    r0[11].merge(t.rows[2].cells[11]); set_cell_text(r0[11], "เฉลี่ย\nคะแนน", 14, True, fill="D9EAF7")
    r0[12].merge(t.rows[2].cells[12]); set_cell_text(r0[12], "จำนวนชั่วโมง\n(ท/ป)", 14, True, fill="D9EAF7")
    # Middle header
    r1 = t.rows[1].cells
    r1[1].merge(r1[6]); set_cell_text(r1[1], "พุทธิพิสัย", 14, True, fill="D9EAF7")
    set_cell_text(r1[7], "ทักษะพิสัย", 14, True, fill="D9EAF7")
    set_cell_text(r1[8], "จิตพิสัย", 14, True, fill="D9EAF7")
    set_cell_text(r1[9], "ประยุกต์ใช้", 14, True, fill="D9EAF7")
    # Bottom header
    r2 = t.rows[2].cells
    for i, h in enumerate(["ความรู้", "ความเข้าใจ", "นำไปใช้", "วิเคราะห์", "ประเมินค่า", "สร้างสรรค์"], start=1):
        set_cell_text(r2[i], h, 14, True, fill="D9EAF7")
    set_repeat_table_header(t.rows[0]); set_repeat_table_header(t.rows[1]); set_repeat_table_header(t.rows[2])

    totals = [0]*6
    total_t = sum(int(x.get("theory", 0) or 0) for x in units)
    total_p = sum(int(x.get("practice", 0) or 0) for x in units)
    for item in units:
        r = t.add_row().cells
        set_cell_text(r[0], f"{item['no']}. {item['title']}", 14, False, WD_ALIGN_PARAGRAPH.LEFT)
        bloom = item.get("bloom", bloom_for_unit(int(item['no'])))
        for i, val in enumerate(bloom, start=1):
            set_cell_text(r[i], val, 14)
            if str(val) != "-":
                try: totals[i-1] += int(val)
                except Exception: pass
        set_cell_text(r[7], "30", 14)
        set_cell_text(r[8], "20", 14)
        set_cell_text(r[9], "10", 14)
        set_cell_text(r[10], "100", 14)
        set_cell_text(r[11], "10", 14)
        set_cell_text(r[12], f"{item.get('theory',0)}/{item.get('practice',0)}", 14)
    r = t.add_row().cells
    set_cell_text(r[0], "รวม", 14, True)
    for i, val in enumerate(totals, start=1):
        set_cell_text(r[i], val if val else 0, 14, True)
    set_cell_text(r[7], "300", 14, True)
    set_cell_text(r[8], "200", 14, True)
    set_cell_text(r[9], "100", 14, True)
    set_cell_text(r[10], "1000", 14, True)
    set_cell_text(r[11], "100", 14, True)
    set_cell_text(r[12], f"{total_t}/{total_p}", 14, True)
    style_table(t, 14, widths)


def add_course_line(doc: Document, course: Course):
    add_p(doc, f"รหัสวิชา {course.code} ชื่อวิชา {course.name}", 16, False, WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, f"ทฤษฎี {course.theory_hours} ชั่วโมง/สัปดาห์  ปฏิบัติ {course.practice_hours} ชั่วโมง/สัปดาห์   จำนวน {course.credits} หน่วยกิต", 16, False, WD_ALIGN_PARAGRAPH.CENTER, after=4)


def add_behavior_table2(doc: Document, course: Course, units: list[dict[str, Any]]):
    add_p(doc, "ตารางวิเคราะห์พฤติกรรมการเรียนรู้", 20, True, WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_course_line(doc, course)
    widths = [12.5, 3.0, 2.3, 2.3, 2.3, 2.7, 2.7]
    t = doc.add_table(rows=2, cols=7)
    r0 = t.rows[0].cells
    r0[0].merge(t.rows[1].cells[0]); set_cell_text(r0[0], "ชื่อหน่วยการเรียนรู้", 14, True, fill="D9EAF7")
    r0[1].merge(r0[4]); set_cell_text(r0[1], "ระดับความสามารถที่คาดหวัง", 14, True, fill="D9EAF7")
    r0[5].merge(t.rows[1].cells[5]); set_cell_text(r0[5], "จำนวนชั่วโมง\n(ท/ป)", 14, True, fill="D9EAF7")
    r0[6].merge(t.rows[1].cells[6]); set_cell_text(r0[6], "ร้อยละ\nประเมินผล", 14, True, fill="D9EAF7")
    for i, h in enumerate(["พุทธิพิสัย\n(K)", "ทักษะพิสัย\n(S)", "จิตพิสัย\n(A)", "ประยุกต์ใช้\n(AP)"], start=1):
        set_cell_text(t.rows[1].cells[i], h, 14, True, fill="D9EAF7")
    set_repeat_table_header(t.rows[0]); set_repeat_table_header(t.rows[1])
    total_t = sum(int(x.get("theory", 0) or 0) for x in units)
    total_p = sum(int(x.get("practice", 0) or 0) for x in units)
    for item in units:
        r = t.add_row().cells
        set_cell_text(r[0], f"หน่วยที่ {item['no']}: {item['title']}", 14, False, WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(r[1], str(item.get("k", k_for_unit(int(item['no'])))), 14)
        set_cell_text(r[2], item.get("s", "S4"), 14)
        set_cell_text(r[3], item.get("a", "A5"), 14)
        set_cell_text(r[4], normalize_ap(item.get("ap", "AP4")), 14)
        set_cell_text(r[5], f"{item.get('theory',0)}/{item.get('practice',0)}", 14)
        set_cell_text(r[6], "10", 14)
    r = t.add_row().cells
    r[0].merge(r[4])
    set_cell_text(r[0], "รวมตลอดภาคเรียน", 14, True, WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(r[5], f"{total_t}/{total_p}", 14, True)
    set_cell_text(r[6], "100", 14, True)
    style_table(t, 14, widths)


def add_unit_table(doc: Document, course: Course, units: list[dict[str, Any]]):
    add_p(doc, "หน่วยการเรียนรู้", 20, True, WD_ALIGN_PARAGRAPH.CENTER, after=6)
    add_course_line(doc, course)
    widths = [2.2, 13.5, 2.8, 2.8, 2.8]
    t = doc.add_table(rows=2, cols=5)
    r0 = t.rows[0].cells
    r0[0].merge(t.rows[1].cells[0]); set_cell_text(r0[0], "หน่วยที่", 16, True, fill="D9EAF7")
    r0[1].merge(t.rows[1].cells[1]); set_cell_text(r0[1], "หน่วยการเรียนรู้", 16, True, fill="D9EAF7")
    r0[2].merge(r0[4]); set_cell_text(r0[2], "เวลาเรียน (ชม.)", 16, True, fill="D9EAF7")
    for i, h in enumerate(["ทฤษฎี", "ปฏิบัติ", "รวม"], start=2):
        set_cell_text(t.rows[1].cells[i], h, 16, True, fill="D9EAF7")
    set_repeat_table_header(t.rows[0]); set_repeat_table_header(t.rows[1])
    total_t = total_p = 0
    for item in units:
        th = int(item.get("theory", 0) or 0)
        pr = int(item.get("practice", 0) or 0)
        total_t += th; total_p += pr
        r = t.add_row().cells
        set_cell_text(r[0], item["no"], 16)
        set_cell_text(r[1], item["title"], 16, False, WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(r[2], th, 16)
        set_cell_text(r[3], pr, 16)
        set_cell_text(r[4], th + pr, 16)
    r = t.add_row().cells
    r[0].merge(r[1])
    set_cell_text(r[0], "รวม", 16, True)
    set_cell_text(r[2], total_t, 16, True)
    set_cell_text(r[3], total_p, 16, True)
    set_cell_text(r[4], total_t + total_p, 16, True)
    style_table(t, 16, widths)


def add_assessment_and_signature(doc: Document, course: Course, teacher: dict[str, Any]):
    add_p(doc, "การประเมินผลลัพธ์การเรียนรู้ระดับรายวิชา", 16, True, WD_ALIGN_PARAGRAPH.LEFT, before=8)
    body = clean_text(course.assessment) or f"ให้ผู้เรียนจัดทำชิ้นงาน/ภาระงานที่สอดคล้องกับรายวิชา {course.name} พร้อมนำเสนอและรับการประเมินตามสมรรถนะรายวิชา"
    add_multiline(doc, body, first_line=1.0)
    add_p(doc, "", 16)
    add_p(doc, "ลงชื่อ.................................................... ผู้จัดทำ", 16, False, WD_ALIGN_PARAGRAPH.RIGHT)
    add_p(doc, f"({teacher.get('teacher_name','')})", 16, False, WD_ALIGN_PARAGRAPH.RIGHT)


def build_teaching_plan(course: Course, teacher: dict[str, Any]) -> Path:
    GENERATED_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _set_doc_defaults(doc)
    set_portrait(doc.sections[0])

    add_cover(doc, course, teacher)
    doc.add_page_break()
    add_course_syllabus(doc, course)

    # All large tables in the model template are landscape.
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    set_landscape(sec)
    add_standard_table(doc, course)
    doc.add_page_break()

    units = build_units_from_description(course)
    add_learning_unit_analysis(doc, course, units)
    doc.add_page_break()
    add_behavior_table1(doc, course, units)
    doc.add_page_break()
    add_behavior_table2(doc, course, units)
    doc.add_page_break()
    add_unit_table(doc, course, units)
    add_assessment_and_signature(doc, course, teacher)

    safe_name = re.sub(r"[\\/:*?\"<>|]+", "_", course.name)[:80]
    filename = f"โครงการสอน_{course.code}_{safe_name}_{uuid.uuid4().hex[:8]}.docx"
    out = GENERATED_DIR / filename
    doc.save(out)
    return out
